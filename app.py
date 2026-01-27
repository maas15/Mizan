"""
Mizan GRC - Enterprise Governance, Risk & Compliance Platform
Flask Application - Full RTL Support
Created by: Eng. Mohammad Abbas Alsaadon
"""

import os
import json
import sqlite3
import hashlib
import secrets
from datetime import datetime, timedelta
from functools import wraps
from flask import Flask, render_template, request, redirect, url_for, session, jsonify, flash
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv('SECRET_KEY', secrets.token_hex(32))
app.permanent_session_lifetime = timedelta(hours=2)

# Configuration
class Config:
    APP_NAME = "Mizan"
    APP_VERSION = "3.0.0"
    APP_TAGLINE = "Governance • Risk • Compliance"
    CREATOR_NAME = "Eng. Mohammad Abbas Alsaadon"
    CREATOR_TITLE = "Consultant/Expert"
    COPYRIGHT_YEAR = "2026"
    DB_PATH = "mizan.db"
    OPENAI_API_KEY = os.getenv('OPENAI_API_KEY', '')

config = Config()

# ============================================================================
# DATABASE
# ============================================================================

def get_db():
    """Get database connection."""
    conn = sqlite3.connect(config.DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    """Initialize database tables."""
    conn = get_db()
    cursor = conn.cursor()
    
    # Users table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Strategies table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS strategies (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            domain TEXT,
            org_name TEXT,
            sector TEXT,
            content TEXT,
            language TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    # Policies table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS policies (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            domain TEXT,
            policy_name TEXT,
            framework TEXT,
            content TEXT,
            language TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    # Risks table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS risks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            domain TEXT,
            asset_name TEXT,
            threat TEXT,
            risk_level TEXT,
            analysis TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    conn.commit()
    conn.close()

# Initialize database on startup
init_db()

# ============================================================================
# AUTHENTICATION
# ============================================================================

def hash_password(password):
    """Hash password with salt."""
    salt = secrets.token_hex(16)
    hash_obj = hashlib.pbkdf2_hmac('sha256', password.encode(), salt.encode(), 100000)
    return f"{salt}${hash_obj.hex()}"

def verify_password(password, stored_hash):
    """Verify password against stored hash."""
    try:
        salt, hash_value = stored_hash.split('$')
        hash_obj = hashlib.pbkdf2_hmac('sha256', password.encode(), salt.encode(), 100000)
        return hash_obj.hex() == hash_value
    except:
        return False

def login_required(f):
    """Decorator to require login."""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

# ============================================================================
# TRANSLATIONS
# ============================================================================

TRANSLATIONS = {
    "en": {
        "app_name": "Mizan",
        "tagline": "Governance • Risk • Compliance",
        "login": "Sign In",
        "register": "Register",
        "logout": "Logout",
        "username": "Username",
        "password": "Password",
        "welcome": "Welcome to Enterprise GRC Platform",
        "welcome_sub": "Secure • Compliant • Intelligent",
        "disclaimer": "AI-generated content. Verify with professionals.",
        "no_sensitive": "Do not enter sensitive or confidential data",
        "domains": ["Cyber Security", "Data Management", "Artificial Intelligence", "Digital Transformation", "Global Standards"],
        "tabs": ["Strategy", "Policy Lab", "Audit", "Risk Radar"],
        "org_name": "Organization Name",
        "sector": "Sector",
        "sectors": ["Government", "Banking/Finance", "Healthcare", "Energy", "Telecom", "Retail", "Manufacturing"],
        "size": "Organization Size",
        "sizes": ["Small (<100)", "Medium (100-1000)", "Large (1000+)"],
        "budget": "Budget Range",
        "budgets": ["< 1M SAR", "1M-5M SAR", "5M-20M SAR", "20M+ SAR"],
        "frameworks": "Regulatory Frameworks",
        "horizon": "Strategic Horizon (Months)",
        "generate": "Generate Strategy",
        "generating": "Generating...",
        "current_state": "Current State Assessment",
        "technologies": "Current Technologies",
        "challenges": "Key Challenges",
        "strategy_sections": {
            "vision": "Executive Vision & Strategic Objectives",
            "gaps": "Current State Assessment (Gap Analysis)",
            "pillars": "Strategic Pillars & Initiatives",
            "roadmap": "Implementation Roadmap",
            "kpis": "Measuring Success (KPIs & KRIs)",
            "confidence": "Confidence Score"
        },
        "policy_name": "Policy Title",
        "policy_framework": "Framework/Standard",
        "generate_policy": "Generate Policy",
        "risk_category": "Risk Category",
        "risk_scenario": "Risk Scenario",
        "asset_name": "Asset Name",
        "analyze_risk": "Analyze Risk",
        "download": "Download",
        "created_by": "Created by",
        "settings": "Settings",
        "clear_history": "Clear History",
        "ai_connected": "AI Core Connected",
        "ai_disconnected": "Simulation Mode",
        "logged_in_as": "Logged in as"
    },
    "ar": {
        "app_name": "ميزان",
        "tagline": "الحوكمة • المخاطر • الامتثال",
        "login": "تسجيل الدخول",
        "register": "إنشاء حساب",
        "logout": "تسجيل الخروج",
        "username": "اسم المستخدم",
        "password": "كلمة المرور",
        "welcome": "مرحباً بك في منصة حوكمة المؤسسات",
        "welcome_sub": "آمن • ممتثل • ذكي",
        "disclaimer": "محتوى مُنشأ بالذكاء الاصطناعي. يُرجى التحقق مع المختصين.",
        "no_sensitive": "لا تدخل بيانات حساسة أو سرية",
        "domains": ["الأمن السيبراني", "إدارة البيانات", "الذكاء الاصطناعي", "التحول الرقمي", "المعايير العالمية"],
        "tabs": ["الاستراتيجية", "معمل السياسات", "التدقيق", "رادار المخاطر"],
        "org_name": "اسم المنظمة",
        "sector": "القطاع",
        "sectors": ["حكومي", "بنوك/مالي", "رعاية صحية", "طاقة", "اتصالات", "تجزئة", "تصنيع"],
        "size": "حجم المنظمة",
        "sizes": ["صغيرة (أقل من 100)", "متوسطة (100-1000)", "كبيرة (أكثر من 1000)"],
        "budget": "نطاق الميزانية",
        "budgets": ["< 1 مليون ريال", "1-5 مليون ريال", "5-20 مليون ريال", "20+ مليون ريال"],
        "frameworks": "الأطر التنظيمية",
        "horizon": "الأفق الاستراتيجي (أشهر)",
        "generate": "إنشاء الاستراتيجية",
        "generating": "جاري الإنشاء...",
        "current_state": "تقييم الوضع الحالي",
        "technologies": "التقنيات الحالية",
        "challenges": "التحديات الرئيسية",
        "strategy_sections": {
            "vision": "الرؤية التنفيذية والأهداف الاستراتيجية",
            "gaps": "تقييم الوضع الراهن (تحليل الفجوات)",
            "pillars": "الركائز الاستراتيجية والمبادرات",
            "roadmap": "خارطة طريق التنفيذ",
            "kpis": "قياس النجاح (مؤشرات الأداء والمخاطر)",
            "confidence": "درجة الثقة والتحقق"
        },
        "policy_name": "عنوان السياسة",
        "policy_framework": "الإطار/المعيار",
        "generate_policy": "إنشاء السياسة",
        "risk_category": "فئة المخاطر",
        "risk_scenario": "سيناريو الخطر",
        "asset_name": "اسم الأصل",
        "analyze_risk": "تحليل الخطر",
        "download": "تحميل",
        "created_by": "تم الإنشاء بواسطة",
        "settings": "الإعدادات",
        "clear_history": "مسح السجل",
        "ai_connected": "الذكاء الاصطناعي متصل",
        "ai_disconnected": "وضع المحاكاة",
        "logged_in_as": "مسجل الدخول كـ"
    }
}

def get_text(lang='en'):
    """Get translations for language."""
    return TRANSLATIONS.get(lang, TRANSLATIONS['en'])

# ============================================================================
# DOMAIN DATA
# ============================================================================

DOMAIN_FRAMEWORKS = {
    "cyber": [
        "NCA ECC (Essential Cybersecurity Controls)",
        "NCA CSCC (Critical Systems Cybersecurity Controls)", 
        "NCA DCC (Data Cybersecurity Controls)",
        "NCA OTCC (Operational Technology Cybersecurity Controls)",
        "NCA TCC (Telework Cybersecurity Controls)",
        "NCA OSMACC (Social Media Cybersecurity Controls)",
        "NCA CCC (Cloud Cybersecurity Controls)",
        "NCA NCS (National Cryptographic Standards)",
        "NCA CGIoT (Cybersecurity Guidelines for IoT)",
        "SAMA CSF",
        "ISO 27001:2022"
    ],
    "data": ["NDMO/SDAIA", "PDPL", "GDPR", "NCA DCC", "DGA Standards"],
    "ai": ["SDAIA AI Ethics", "NIST AI RMF", "EU AI Act", "ISO 42001"],
    "dt": ["DGA Digital Policy", "COBIT 2019", "TOGAF", "ITIL 4"],
    "global": ["ISO 27001:2022", "ISO 22301", "NIST CSF 2.0", "ISO 9001", "ISO 31000"]
}

# Domain-specific foundational technologies/controls
DOMAIN_TECHNOLOGIES = {
    "cyber": {
        "en": {
            "Security Operations": ["SIEM", "SOAR", "SOC", "Threat Intelligence Platform", "Log Management"],
            "Endpoint Security": ["EDR/XDR", "Antivirus/Anti-malware", "Mobile Device Management (MDM)", "Endpoint DLP"],
            "Network Security": ["Next-Gen Firewall", "IDS/IPS", "Network Access Control (NAC)", "Web Proxy", "DNS Security"],
            "Identity & Access": ["IAM", "PAM", "MFA/2FA", "SSO", "Directory Services (AD/LDAP)"],
            "Data Protection": ["DLP", "Encryption (at-rest/in-transit)", "Backup & Recovery", "Data Classification"],
            "Application Security": ["WAF", "SAST/DAST", "API Security", "Code Review Tools"],
            "Cloud Security": ["CASB", "CSPM", "CWPP", "Cloud IAM"],
            "GRC Tools": ["Vulnerability Scanner", "Penetration Testing", "Compliance Management", "Risk Register"]
        },
        "ar": {
            "العمليات الأمنية": ["SIEM", "SOAR", "مركز العمليات الأمنية", "منصة استخبارات التهديدات", "إدارة السجلات"],
            "أمن النقاط الطرفية": ["EDR/XDR", "مكافحة الفيروسات", "إدارة الأجهزة المحمولة", "DLP للنقاط الطرفية"],
            "أمن الشبكات": ["جدار الحماية المتقدم", "IDS/IPS", "التحكم بالوصول للشبكة", "بروكسي الويب", "أمن DNS"],
            "الهوية والوصول": ["IAM", "PAM", "المصادقة متعددة العوامل", "SSO", "خدمات الدليل"],
            "حماية البيانات": ["DLP", "التشفير", "النسخ الاحتياطي والاستعادة", "تصنيف البيانات"],
            "أمن التطبيقات": ["WAF", "SAST/DAST", "أمن API", "أدوات مراجعة الكود"],
            "أمن السحابة": ["CASB", "CSPM", "CWPP", "IAM السحابي"],
            "أدوات الحوكمة": ["ماسح الثغرات", "اختبار الاختراق", "إدارة الامتثال", "سجل المخاطر"]
        }
    },
    "data": {
        "en": {
            "Data Governance": ["Data Catalog", "Metadata Management", "Data Lineage", "Business Glossary"],
            "Data Quality": ["Data Profiling", "Data Cleansing", "Master Data Management (MDM)", "Data Validation"],
            "Data Security": ["Data Masking", "Tokenization", "Database Encryption", "Access Controls"],
            "Data Privacy": ["Consent Management", "Privacy Impact Assessment", "Data Subject Rights Management", "Cookie Management"],
            "Data Integration": ["ETL/ELT Tools", "Data Virtualization", "API Management", "Data Replication"],
            "Data Analytics": ["BI Platform", "Data Warehouse", "Data Lake", "Reporting Tools"],
            "Data Lifecycle": ["Archiving Solutions", "Retention Management", "Secure Disposal", "Backup Systems"]
        },
        "ar": {
            "حوكمة البيانات": ["كتالوج البيانات", "إدارة البيانات الوصفية", "تتبع مسار البيانات", "قاموس الأعمال"],
            "جودة البيانات": ["تحليل البيانات", "تنظيف البيانات", "إدارة البيانات الرئيسية", "التحقق من البيانات"],
            "أمن البيانات": ["إخفاء البيانات", "الترميز", "تشفير قواعد البيانات", "ضوابط الوصول"],
            "خصوصية البيانات": ["إدارة الموافقات", "تقييم أثر الخصوصية", "إدارة حقوق أصحاب البيانات", "إدارة الكوكيز"],
            "تكامل البيانات": ["أدوات ETL/ELT", "المحاكاة الافتراضية للبيانات", "إدارة API", "نسخ البيانات"],
            "تحليل البيانات": ["منصة ذكاء الأعمال", "مستودع البيانات", "بحيرة البيانات", "أدوات التقارير"],
            "دورة حياة البيانات": ["حلول الأرشفة", "إدارة الاحتفاظ", "الإتلاف الآمن", "أنظمة النسخ الاحتياطي"]
        }
    },
    "ai": {
        "en": {
            "ML Infrastructure": ["ML Platform", "Model Registry", "Feature Store", "Experiment Tracking"],
            "Data for AI": ["Data Labeling", "Training Data Management", "Synthetic Data Generation", "Data Versioning"],
            "Model Development": ["AutoML", "Notebook Environment", "Model Training Pipeline", "Hyperparameter Tuning"],
            "Model Operations": ["Model Deployment", "Model Monitoring", "A/B Testing", "Model Versioning"],
            "AI Governance": ["Model Documentation", "Bias Detection", "Explainability Tools", "Audit Trail"],
            "AI Security": ["Adversarial Testing", "Model Encryption", "Secure Inference", "Access Controls"]
        },
        "ar": {
            "بنية تعلم الآلة": ["منصة ML", "سجل النماذج", "مخزن الميزات", "تتبع التجارب"],
            "بيانات الذكاء الاصطناعي": ["تصنيف البيانات", "إدارة بيانات التدريب", "توليد البيانات الاصطناعية", "إصدارات البيانات"],
            "تطوير النماذج": ["AutoML", "بيئة Notebook", "خط أنابيب التدريب", "ضبط المعاملات"],
            "عمليات النماذج": ["نشر النماذج", "مراقبة النماذج", "اختبار A/B", "إصدارات النماذج"],
            "حوكمة الذكاء الاصطناعي": ["توثيق النماذج", "كشف التحيز", "أدوات التفسير", "سجل المراجعة"],
            "أمن الذكاء الاصطناعي": ["الاختبار العدائي", "تشفير النماذج", "الاستدلال الآمن", "ضوابط الوصول"]
        }
    },
    "dt": {
        "en": {
            "Digital Platforms": ["Enterprise Portal", "Mobile Apps", "Customer Experience Platform", "Digital Workplace"],
            "Integration": ["ESB/Integration Platform", "API Gateway", "iPaaS", "Microservices"],
            "Process Automation": ["RPA", "BPM", "Workflow Engine", "Low-Code Platform"],
            "Cloud Services": ["IaaS", "PaaS", "SaaS", "Hybrid Cloud"],
            "Analytics & Insights": ["Big Data Platform", "Real-time Analytics", "Predictive Analytics", "Dashboard/KPI Tools"],
            "Collaboration": ["Unified Communications", "Document Management", "Project Management", "Knowledge Management"],
            "Customer Engagement": ["CRM", "Marketing Automation", "Chatbots", "Omnichannel Platform"]
        },
        "ar": {
            "المنصات الرقمية": ["البوابة المؤسسية", "تطبيقات الجوال", "منصة تجربة العميل", "مكان العمل الرقمي"],
            "التكامل": ["منصة التكامل", "بوابة API", "iPaaS", "الخدمات المصغرة"],
            "أتمتة العمليات": ["RPA", "إدارة العمليات", "محرك سير العمل", "منصة Low-Code"],
            "الخدمات السحابية": ["IaaS", "PaaS", "SaaS", "السحابة الهجينة"],
            "التحليلات": ["منصة البيانات الضخمة", "التحليلات الفورية", "التحليلات التنبؤية", "لوحات المؤشرات"],
            "التعاون": ["الاتصالات الموحدة", "إدارة الوثائق", "إدارة المشاريع", "إدارة المعرفة"],
            "تفاعل العملاء": ["CRM", "أتمتة التسويق", "روبوتات المحادثة", "منصة القنوات المتعددة"]
        }
    },
    "global": {
        "en": {
            "Quality Management": ["QMS Software", "Document Control", "CAPA Management", "Audit Management"],
            "Risk Management": ["ERM Platform", "Risk Register", "Risk Assessment Tools", "Incident Management"],
            "Compliance": ["Compliance Management", "Policy Management", "Training Management", "Certification Tracking"],
            "Business Continuity": ["BCP Platform", "DR Solutions", "Crisis Management", "Emergency Notification"],
            "Information Security": ["ISMS Platform", "Asset Management", "Vulnerability Management", "Security Awareness"]
        },
        "ar": {
            "إدارة الجودة": ["برنامج QMS", "التحكم بالوثائق", "إدارة CAPA", "إدارة التدقيق"],
            "إدارة المخاطر": ["منصة ERM", "سجل المخاطر", "أدوات تقييم المخاطر", "إدارة الحوادث"],
            "الامتثال": ["إدارة الامتثال", "إدارة السياسات", "إدارة التدريب", "تتبع الشهادات"],
            "استمرارية الأعمال": ["منصة BCP", "حلول DR", "إدارة الأزمات", "الإشعارات الطارئة"],
            "أمن المعلومات": ["منصة ISMS", "إدارة الأصول", "إدارة الثغرات", "التوعية الأمنية"]
        }
    }
}

# Enhanced Risk Categories with scenarios
RISK_CATEGORIES = {
    "cyber": {
        "en": {
            "Access Control": ["Unauthorized access to systems", "Privilege escalation", "Credential theft", "Insider threat", "Session hijacking"],
            "Network Security": ["Network intrusion", "DDoS attack", "Man-in-the-middle attack", "DNS poisoning", "Lateral movement"],
            "Data Protection": ["Data breach", "Data leakage", "Unauthorized data access", "Data corruption", "Ransomware encryption"],
            "Endpoint Security": ["Malware infection", "Zero-day exploit", "USB-based attack", "Remote access trojan", "Cryptomining"],
            "Application Security": ["SQL injection", "XSS attack", "API abuse", "Broken authentication", "Insecure deserialization"],
            "Cloud Security": ["Cloud misconfiguration", "Data exposure in cloud", "Account hijacking", "Insecure APIs", "Shadow IT"],
            "Social Engineering": ["Phishing attack", "Spear phishing", "Business email compromise", "Vishing", "Pretexting"],
            "Third Party Risk": ["Vendor breach", "Supply chain attack", "Third-party data exposure", "Service provider failure"],
            "Operational Technology": ["SCADA attack", "ICS compromise", "Physical-cyber attack", "OT network breach"],
            "Incident Response": ["Delayed detection", "Inadequate response", "Evidence loss", "Communication failure"]
        },
        "ar": {
            "التحكم بالوصول": ["وصول غير مصرح للأنظمة", "تصعيد الصلاحيات", "سرقة بيانات الاعتماد", "التهديد الداخلي", "اختطاف الجلسة"],
            "أمن الشبكات": ["اختراق الشبكة", "هجوم DDoS", "هجوم الوسيط", "تسميم DNS", "الحركة الجانبية"],
            "حماية البيانات": ["خرق البيانات", "تسرب البيانات", "وصول غير مصرح للبيانات", "تلف البيانات", "تشفير الفدية"],
            "أمن النقاط الطرفية": ["إصابة بالبرمجيات الخبيثة", "استغلال يوم الصفر", "هجوم USB", "حصان طروادة", "التعدين الخبيث"],
            "أمن التطبيقات": ["حقن SQL", "هجوم XSS", "إساءة استخدام API", "مصادقة معطلة", "إلغاء تسلسل غير آمن"],
            "أمن السحابة": ["سوء تكوين السحابة", "كشف البيانات السحابية", "اختطاف الحساب", "APIs غير آمنة", "Shadow IT"],
            "الهندسة الاجتماعية": ["هجوم التصيد", "التصيد الموجه", "اختراق البريد التجاري", "التصيد الصوتي", "الذريعة"],
            "مخاطر الأطراف الثالثة": ["اختراق المورد", "هجوم سلسلة التوريد", "كشف بيانات الطرف الثالث", "فشل مزود الخدمة"],
            "التقنيات التشغيلية": ["هجوم SCADA", "اختراق ICS", "هجوم فيزيائي-سيبراني", "اختراق شبكة OT"],
            "الاستجابة للحوادث": ["تأخر الكشف", "استجابة غير كافية", "فقدان الأدلة", "فشل الاتصال"]
        }
    },
    "data": {
        "en": {
            "Data Quality": ["Incomplete data", "Duplicate records", "Data inconsistency", "Outdated information", "Invalid data formats"],
            "Data Privacy": ["Personal data exposure", "Consent violation", "Cross-border transfer issues", "Right to erasure failure", "Purpose limitation breach"],
            "Data Governance": ["Undefined data ownership", "Missing data lineage", "Inconsistent data definitions", "Policy non-compliance", "Metadata gaps"],
            "Data Security": ["Unauthorized data access", "Database breach", "Encryption failure", "Backup exposure", "Insider data theft"],
            "Data Lifecycle": ["Retention policy violation", "Improper disposal", "Archive corruption", "Recovery failure", "Storage overflow"],
            "Data Integration": ["ETL failure", "Data sync issues", "API data exposure", "Migration errors", "Real-time feed disruption"],
            "Regulatory Compliance": ["PDPL violation", "GDPR non-compliance", "Audit findings", "Reporting failures", "Documentation gaps"]
        },
        "ar": {
            "جودة البيانات": ["بيانات ناقصة", "سجلات مكررة", "عدم اتساق البيانات", "معلومات قديمة", "تنسيقات بيانات غير صالحة"],
            "خصوصية البيانات": ["كشف البيانات الشخصية", "انتهاك الموافقة", "مشاكل النقل عبر الحدود", "فشل حق المحو", "انتهاك تحديد الغرض"],
            "حوكمة البيانات": ["ملكية بيانات غير محددة", "مسار بيانات مفقود", "تعريفات بيانات غير متسقة", "عدم الامتثال للسياسة", "فجوات البيانات الوصفية"],
            "أمن البيانات": ["وصول غير مصرح للبيانات", "اختراق قاعدة البيانات", "فشل التشفير", "كشف النسخ الاحتياطية", "سرقة البيانات الداخلية"],
            "دورة حياة البيانات": ["انتهاك سياسة الاحتفاظ", "التخلص غير السليم", "تلف الأرشيف", "فشل الاستعادة", "تجاوز التخزين"],
            "تكامل البيانات": ["فشل ETL", "مشاكل مزامنة البيانات", "كشف بيانات API", "أخطاء الترحيل", "انقطاع التغذية الفورية"],
            "الامتثال التنظيمي": ["انتهاك PDPL", "عدم الامتثال لـ GDPR", "نتائج التدقيق", "فشل التقارير", "فجوات التوثيق"]
        }
    },
    "ai": {
        "en": {
            "Model Bias": ["Demographic bias", "Selection bias", "Measurement bias", "Algorithmic discrimination", "Feedback loop bias"],
            "Data Quality for AI": ["Training data poisoning", "Label errors", "Data drift", "Insufficient training data", "Unrepresentative samples"],
            "Model Performance": ["Model degradation", "Concept drift", "Overfitting", "Underfitting", "Poor generalization"],
            "Explainability": ["Black box decisions", "Lack of interpretability", "Unexplainable outcomes", "Audit trail gaps", "Stakeholder confusion"],
            "AI Security": ["Adversarial attacks", "Model theft", "Model inversion", "Data extraction attacks", "Prompt injection"],
            "Ethical Concerns": ["Autonomous harm", "Privacy invasion", "Unfair treatment", "Manipulation", "Job displacement"],
            "Operational Risks": ["Model failure in production", "Integration issues", "Scalability problems", "Resource constraints", "Dependency failures"],
            "Compliance": ["Regulatory non-compliance", "Documentation gaps", "Consent issues", "Cross-border AI use", "Audit failures"]
        },
        "ar": {
            "تحيز النموذج": ["تحيز ديموغرافي", "تحيز الاختيار", "تحيز القياس", "تمييز خوارزمي", "تحيز حلقة التغذية"],
            "جودة بيانات الذكاء الاصطناعي": ["تسميم بيانات التدريب", "أخطاء التصنيف", "انحراف البيانات", "بيانات تدريب غير كافية", "عينات غير ممثلة"],
            "أداء النموذج": ["تدهور النموذج", "انحراف المفهوم", "الإفراط في التخصيص", "نقص التخصيص", "ضعف التعميم"],
            "قابلية التفسير": ["قرارات الصندوق الأسود", "نقص قابلية التفسير", "نتائج غير قابلة للتفسير", "فجوات سجل المراجعة", "إرباك أصحاب المصلحة"],
            "أمن الذكاء الاصطناعي": ["الهجمات العدائية", "سرقة النموذج", "عكس النموذج", "هجمات استخراج البيانات", "حقن الأوامر"],
            "المخاوف الأخلاقية": ["الضرر المستقل", "انتهاك الخصوصية", "المعاملة غير العادلة", "التلاعب", "استبدال الوظائف"],
            "المخاطر التشغيلية": ["فشل النموذج في الإنتاج", "مشاكل التكامل", "مشاكل قابلية التوسع", "قيود الموارد", "فشل التبعيات"],
            "الامتثال": ["عدم الامتثال التنظيمي", "فجوات التوثيق", "مشاكل الموافقة", "استخدام AI عبر الحدود", "فشل التدقيق"]
        }
    },
    "dt": {
        "en": {
            "Change Management": ["Resistance to change", "Inadequate training", "Cultural barriers", "Communication gaps", "Leadership misalignment"],
            "Technology Integration": ["System incompatibility", "Data migration failures", "API integration issues", "Legacy system constraints", "Vendor lock-in"],
            "Digital Skills": ["Skills shortage", "Knowledge gaps", "Training inadequacy", "Talent retention", "Digital literacy"],
            "Process Disruption": ["Workflow disruption", "Process automation failure", "Business continuity impact", "Operational inefficiency", "Service degradation"],
            "Customer Experience": ["Poor digital experience", "Channel inconsistency", "Accessibility issues", "Response time delays", "Customer data exposure"],
            "Strategic Alignment": ["Misaligned objectives", "ROI uncertainty", "Scope creep", "Priority conflicts", "Resource constraints"],
            "Vendor & Cloud": ["Vendor dependency", "Cloud service outage", "Contract issues", "Cost overruns", "SLA breaches"],
            "Security in Transformation": ["Security gaps during migration", "New attack surfaces", "Access control issues", "Data exposure", "Compliance gaps"]
        },
        "ar": {
            "إدارة التغيير": ["مقاومة التغيير", "تدريب غير كافٍ", "حواجز ثقافية", "فجوات الاتصال", "عدم توافق القيادة"],
            "تكامل التقنية": ["عدم توافق الأنظمة", "فشل ترحيل البيانات", "مشاكل تكامل API", "قيود الأنظمة القديمة", "الارتباط بالمورد"],
            "المهارات الرقمية": ["نقص المهارات", "فجوات المعرفة", "قصور التدريب", "الاحتفاظ بالمواهب", "الثقافة الرقمية"],
            "تعطيل العمليات": ["تعطيل سير العمل", "فشل أتمتة العمليات", "تأثير استمرارية الأعمال", "عدم كفاءة التشغيل", "تدهور الخدمة"],
            "تجربة العميل": ["تجربة رقمية سيئة", "عدم اتساق القنوات", "مشاكل إمكانية الوصول", "تأخر وقت الاستجابة", "كشف بيانات العميل"],
            "التوافق الاستراتيجي": ["أهداف غير متوافقة", "عدم يقين العائد", "زحف النطاق", "تعارض الأولويات", "قيود الموارد"],
            "المورد والسحابة": ["الاعتماد على المورد", "انقطاع الخدمة السحابية", "مشاكل العقود", "تجاوز التكاليف", "انتهاكات SLA"],
            "الأمن في التحول": ["فجوات أمنية أثناء الترحيل", "أسطح هجوم جديدة", "مشاكل التحكم بالوصول", "كشف البيانات", "فجوات الامتثال"]
        }
    },
    "global": {
        "en": {
            "Strategic Risk": ["Market changes", "Competitive disruption", "Regulatory changes", "Technology obsolescence", "Geopolitical factors"],
            "Operational Risk": ["Process failures", "System outages", "Human errors", "Resource constraints", "Supplier disruptions"],
            "Financial Risk": ["Budget overruns", "Cost escalation", "Revenue impact", "Investment loss", "Currency fluctuation"],
            "Compliance Risk": ["Regulatory violations", "Audit findings", "Certification loss", "Legal penalties", "Reporting failures"],
            "Reputational Risk": ["Brand damage", "Customer trust loss", "Media exposure", "Stakeholder concerns", "Social media crisis"],
            "Business Continuity": ["Natural disasters", "Pandemic impact", "Infrastructure failure", "Key person dependency", "Supply chain disruption"],
            "Information Security": ["Data breaches", "Cyber attacks", "Insider threats", "Third-party risks", "Physical security"],
            "Quality Risk": ["Product defects", "Service failures", "Customer complaints", "Non-conformance", "Continuous improvement gaps"]
        },
        "ar": {
            "المخاطر الاستراتيجية": ["تغيرات السوق", "التعطيل التنافسي", "التغييرات التنظيمية", "تقادم التقنية", "العوامل الجيوسياسية"],
            "المخاطر التشغيلية": ["فشل العمليات", "انقطاع الأنظمة", "الأخطاء البشرية", "قيود الموارد", "اضطرابات الموردين"],
            "المخاطر المالية": ["تجاوز الميزانية", "تصاعد التكاليف", "تأثير الإيرادات", "خسارة الاستثمار", "تقلب العملة"],
            "مخاطر الامتثال": ["انتهاكات تنظيمية", "نتائج التدقيق", "فقدان الشهادات", "العقوبات القانونية", "فشل التقارير"],
            "مخاطر السمعة": ["ضرر العلامة التجارية", "فقدان ثقة العميل", "التعرض الإعلامي", "مخاوف أصحاب المصلحة", "أزمة وسائل التواصل"],
            "استمرارية الأعمال": ["الكوارث الطبيعية", "تأثير الجائحة", "فشل البنية التحتية", "الاعتماد على أشخاص رئيسيين", "تعطل سلسلة التوريد"],
            "أمن المعلومات": ["خروقات البيانات", "الهجمات السيبرانية", "التهديدات الداخلية", "مخاطر الأطراف الثالثة", "الأمن المادي"],
            "مخاطر الجودة": ["عيوب المنتج", "فشل الخدمة", "شكاوى العملاء", "عدم المطابقة", "فجوات التحسين المستمر"]
        }
    }
}

DOMAIN_CODES = {
    "Cyber Security": "cyber",
    "Data Management": "data", 
    "Artificial Intelligence": "ai",
    "Digital Transformation": "dt",
    "Global Standards": "global",
    "الأمن السيبراني": "cyber",
    "إدارة البيانات": "data",
    "الذكاء الاصطناعي": "ai",
    "التحول الرقمي": "dt",
    "المعايير العالمية": "global"
}

# ============================================================================
# AI SERVICE
# ============================================================================

def check_ai_available():
    """Check if OpenAI API is available."""
    return bool(config.OPENAI_API_KEY)

def generate_ai_content(prompt, language='en'):
    """Generate content using OpenAI API."""
    if not config.OPENAI_API_KEY:
        return generate_simulation_content(prompt, language)
    
    try:
        import openai
        client = openai.OpenAI(api_key=config.OPENAI_API_KEY)
        
        system_prompt = "You are an expert GRC consultant. Provide professional, detailed responses."
        if language == 'ar':
            system_prompt = "أنت مستشار خبير في الحوكمة والمخاطر والامتثال. قدم ردوداً مهنية ومفصلة باللغة العربية."
        
        response = client.chat.completions.create(
            model="gpt-4-turbo",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt}
            ],
            max_tokens=4000,
            temperature=0.7
        )
        
        return response.choices[0].message.content
    except Exception as e:
        print(f"AI Error: {e}")
        return generate_simulation_content(prompt, language)

def generate_simulation_content(prompt, language='en'):
    """Generate simulated content when AI is unavailable."""
    if language == 'ar':
        return """## الرؤية الاستراتيجية
هذا محتوى تجريبي تم إنشاؤه في وضع المحاكاة.

## تحليل الفجوات
يتطلب التحليل الكامل اتصال الذكاء الاصطناعي.

## الركائز الاستراتيجية
1. تعزيز الحوكمة
2. إدارة المخاطر
3. ضمان الامتثال

## خارطة الطريق
المرحلة 1: التقييم (3 أشهر)
المرحلة 2: التنفيذ (6 أشهر)
المرحلة 3: التحسين المستمر

## مؤشرات الأداء
- نسبة الامتثال
- وقت الاستجابة للحوادث
- تغطية التدريب"""
    else:
        return """## Strategic Vision
This is simulated content generated in simulation mode.

## Gap Analysis
Full analysis requires AI connection.

## Strategic Pillars
1. Governance Enhancement
2. Risk Management
3. Compliance Assurance

## Implementation Roadmap
Phase 1: Assessment (3 months)
Phase 2: Implementation (6 months)
Phase 3: Continuous Improvement

## Key Performance Indicators
- Compliance Rate
- Incident Response Time
- Training Coverage"""

# ============================================================================
# ROUTES - AUTHENTICATION
# ============================================================================

@app.route('/')
def index():
    """Home page - redirect to login or dashboard."""
    if 'user_id' in session:
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    """Login page."""
    lang = request.args.get('lang', session.get('lang', 'en'))
    session['lang'] = lang
    txt = get_text(lang)
    
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '')
        
        conn = get_db()
        user = conn.execute('SELECT * FROM users WHERE username = ?', (username,)).fetchone()
        conn.close()
        
        if user and verify_password(password, user['password_hash']):
            session.permanent = True
            session['user_id'] = user['id']
            session['username'] = user['username']
            return redirect(url_for('dashboard'))
        else:
            flash('Invalid username or password', 'error')
    
    return render_template('login.html', 
                          txt=txt, 
                          lang=lang, 
                          config=config,
                          is_rtl=(lang == 'ar'))

@app.route('/register', methods=['POST'])
def register():
    """Register new user."""
    lang = session.get('lang', 'en')
    username = request.form.get('username', '').strip()
    password = request.form.get('password', '')
    
    if len(username) < 3 or len(password) < 6:
        flash('Username must be 3+ chars, password 6+ chars', 'error')
        return redirect(url_for('login', lang=lang))
    
    conn = get_db()
    try:
        conn.execute('INSERT INTO users (username, password_hash) VALUES (?, ?)',
                    (username, hash_password(password)))
        conn.commit()
        flash('Account created! Please login.', 'success')
    except sqlite3.IntegrityError:
        flash('Username already exists', 'error')
    finally:
        conn.close()
    
    return redirect(url_for('login', lang=lang))

@app.route('/logout')
def logout():
    """Logout user."""
    lang = session.get('lang', 'en')
    session.clear()
    session['lang'] = lang
    return redirect(url_for('login', lang=lang))

# ============================================================================
# ROUTES - DASHBOARD
# ============================================================================

@app.route('/dashboard')
@login_required
def dashboard():
    """Main dashboard."""
    lang = request.args.get('lang', session.get('lang', 'en'))
    session['lang'] = lang
    txt = get_text(lang)
    
    # Get user stats
    conn = get_db()
    strategies_count = conn.execute(
        'SELECT COUNT(*) FROM strategies WHERE user_id = ?', 
        (session['user_id'],)
    ).fetchone()[0]
    policies_count = conn.execute(
        'SELECT COUNT(*) FROM policies WHERE user_id = ?',
        (session['user_id'],)
    ).fetchone()[0]
    risks_count = conn.execute(
        'SELECT COUNT(*) FROM risks WHERE user_id = ?',
        (session['user_id'],)
    ).fetchone()[0]
    conn.close()
    
    return render_template('dashboard.html',
                          txt=txt,
                          lang=lang,
                          config=config,
                          is_rtl=(lang == 'ar'),
                          username=session.get('username'),
                          ai_available=check_ai_available(),
                          stats={
                              'strategies': strategies_count,
                              'policies': policies_count,
                              'risks': risks_count
                          },
                          domains=txt['domains'],
                          domain_codes=DOMAIN_CODES,
                          frameworks=DOMAIN_FRAMEWORKS)

@app.route('/domain/<domain_name>')
@login_required
def domain_page(domain_name):
    """Domain-specific page."""
    lang = request.args.get('lang', session.get('lang', 'en'))
    session['lang'] = lang
    txt = get_text(lang)
    
    domain_code = DOMAIN_CODES.get(domain_name, 'global')
    frameworks = DOMAIN_FRAMEWORKS.get(domain_code, [])
    
    # Get domain-specific technologies
    lang_key = 'ar' if lang == 'ar' else 'en'
    technologies = DOMAIN_TECHNOLOGIES.get(domain_code, {}).get(lang_key, {})
    
    # Get risk categories with scenarios
    risk_data = RISK_CATEGORIES.get(domain_code, {}).get(lang_key, {})
    
    return render_template('domain.html',
                          txt=txt,
                          lang=lang,
                          config=config,
                          is_rtl=(lang == 'ar'),
                          username=session.get('username'),
                          ai_available=check_ai_available(),
                          domain_name=domain_name,
                          domain_code=domain_code,
                          frameworks=frameworks,
                          technologies=technologies,
                          risk_categories=risk_data)

# ============================================================================
# ROUTES - API ENDPOINTS
# ============================================================================

@app.route('/api/generate-strategy', methods=['POST'])
@login_required
def api_generate_strategy():
    """Generate strategy via AI."""
    try:
        data = request.json
        lang = data.get('language', 'en')
        
        # Get current state info
        org_structure = data.get('org_structure', 'Not specified')
        technologies = data.get('technologies', [])
        maturity = data.get('maturity_level', 'initial')
        tech_list = ', '.join(technologies) if technologies else 'None specified'
        frameworks_list = ', '.join(data.get('frameworks', [])) if data.get('frameworks') else 'Not specified'
        
        if lang == 'ar':
            prompt = f"""أنت خبير في الحوكمة والمخاطر والامتثال. أنشئ استراتيجية شاملة.

معلومات المنظمة:
- الاسم: {data.get('org_name', 'المنظمة')}
- القطاع: {data.get('sector', 'عام')}
- المجال: {data.get('domain', 'الأمن السيبراني')}
- الحجم: {data.get('size', 'متوسط')}
- الميزانية: {data.get('budget', '1-5 مليون')}
- الأطر التنظيمية: {frameworks_list}
- الهيكل الحالي: {org_structure}
- التقنيات المطبقة: {tech_list}
- مستوى النضج: {maturity}
- التحديات: {data.get('challenges', 'غير محدد')}

اكتب 6 أقسام منفصلة. استخدم [SECTION] كفاصل بين كل قسم:

القسم 1 - الرؤية والأهداف:
رؤية واضحة و 5-7 أهداف محددة

[SECTION]

القسم 2 - تحليل الفجوات:
الفجوات بين الوضع الحالي والمتطلبات

[SECTION]

القسم 3 - الركائز الاستراتيجية:
4-5 ركائز مع مبادرات

[SECTION]

القسم 4 - خارطة الطريق:
المرحلة 1 (0-6 أشهر)، المرحلة 2 (6-12 شهر)، المرحلة 3 (12-24 شهر)

[SECTION]

القسم 5 - مؤشرات الأداء:
8-10 مؤشرات قابلة للقياس

[SECTION]

القسم 6 - تقييم الثقة:
درجة الثقة والمخاطر الرئيسية"""
        else:
            prompt = f"""You are a GRC expert. Generate a comprehensive strategy.

Organization Info:
- Name: {data.get('org_name', 'Organization')}
- Sector: {data.get('sector', 'General')}
- Domain: {data.get('domain', 'Cyber Security')}
- Size: {data.get('size', 'Medium')}
- Budget: {data.get('budget', '1M-5M')}
- Frameworks: {frameworks_list}
- Current Structure: {org_structure}
- Technologies: {tech_list}
- Maturity: {maturity}
- Challenges: {data.get('challenges', 'Not specified')}

Write 6 separate sections. Use [SECTION] as separator between each:

Section 1 - Vision & Objectives:
Clear vision and 5-7 specific objectives

[SECTION]

Section 2 - Gap Analysis:
Gaps between current state and requirements

[SECTION]

Section 3 - Strategic Pillars:
4-5 pillars with initiatives

[SECTION]

Section 4 - Implementation Roadmap:
Phase 1 (0-6 months), Phase 2 (6-12 months), Phase 3 (12-24 months)

[SECTION]

Section 5 - KPIs:
8-10 measurable KPIs

[SECTION]

Section 6 - Confidence Assessment:
Confidence score and key risks"""

        content = generate_ai_content(prompt, lang)
        
        # Parse sections - try multiple separators
        # The AI might use [SECTION], ---, or numbered sections
        parts = []
        
        if '[SECTION]' in content:
            parts = content.split('[SECTION]')
        elif '\n---\n' in content:
            parts = content.split('\n---\n')
        elif '---' in content:
            parts = content.split('---')
        else:
            # Try to split by section headers
            import re
            parts = re.split(r'\n(?=(?:Section \d|القسم \d|##))', content)
        
        # Clean each section
        def clean_section(text):
            if not text:
                return ''
            text = text.strip()
            lines = text.split('\n')
            # Remove section headers if present
            while lines and any(x in lines[0].lower() for x in ['section', 'القسم', '---', '**section', '1.', '2.', '3.', '4.', '5.', '6.']):
                lines = lines[1:]
            # Remove leading dashes or numbers
            result = '\n'.join(lines).strip()
            # Remove leading "1 -", "2 -" etc
            result = re.sub(r'^\d+\s*[-–]\s*', '', result)
            return result
        
        # Ensure we have 6 parts
        while len(parts) < 6:
            parts.append('')
        
        sections = {
            'vision': clean_section(parts[0]),
            'gaps': clean_section(parts[1]),
            'pillars': clean_section(parts[2]),
            'roadmap': clean_section(parts[3]),
            'kpis': clean_section(parts[4]),
            'confidence': clean_section(parts[5]) if len(parts) > 5 else ''
        }
        
        # Save to database
        try:
            conn = get_db()
            conn.execute('''INSERT INTO strategies (user_id, domain, org_name, sector, content, language)
                            VALUES (?, ?, ?, ?, ?, ?)''',
                        (session['user_id'], data.get('domain'), data.get('org_name'), 
                         data.get('sector'), content, lang))
            conn.commit()
            conn.close()
        except Exception as db_error:
            print(f"Database error: {db_error}")
        
        return jsonify({
            'success': True,
            'sections': sections
        })
        
    except Exception as e:
        print(f"Strategy generation error: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500
    
@app.route('/api/generate-policy', methods=['POST'])
@login_required
def api_generate_policy():
    """Generate policy document."""
    data = request.json
    lang = data.get('language', 'en')
    
    prompt = f"""Generate a professional {data.get('policy_name', 'Security')} Policy document based on {data.get('framework', 'ISO 27001')}.
Include: Purpose, Scope, Policy Statements, Roles & Responsibilities, Compliance Requirements, Review Process."""

    if lang == 'ar':
        prompt = f"""أنشئ وثيقة سياسة {data.get('policy_name', 'أمنية')} احترافية بناءً على {data.get('framework', 'ISO 27001')}.
تتضمن: الغرض، النطاق، بنود السياسة، الأدوار والمسؤوليات، متطلبات الامتثال، عملية المراجعة."""

    content = generate_ai_content(prompt, lang)
    
    # Save to database
    conn = get_db()
    conn.execute('''INSERT INTO policies (user_id, domain, policy_name, framework, content, language)
                    VALUES (?, ?, ?, ?, ?, ?)''',
                (session['user_id'], data.get('domain'), data.get('policy_name'),
                 data.get('framework'), content, lang))
    conn.commit()
    conn.close()
    
    return jsonify({'success': True, 'content': content})

@app.route('/api/analyze-risk', methods=['POST'])
@login_required
def api_analyze_risk():
    """Analyze risk scenario."""
    data = request.json
    lang = data.get('language', 'en')
    
    prompt = f"""Analyze this risk scenario:
Category: {data.get('category', 'General')}
Asset: {data.get('asset', 'System')}
Threat: {data.get('threat', 'Unauthorized Access')}

Provide: Risk Assessment, Impact Analysis, Likelihood, Recommended Controls, Residual Risk."""

    if lang == 'ar':
        prompt = f"""حلل سيناريو الخطر التالي:
الفئة: {data.get('category', 'عام')}
الأصل: {data.get('asset', 'النظام')}
التهديد: {data.get('threat', 'وصول غير مصرح')}

قدم: تقييم الخطر، تحليل الأثر، الاحتمالية، الضوابط الموصى بها، الخطر المتبقي."""

    content = generate_ai_content(prompt, lang)
    
    # Save to database
    conn = get_db()
    conn.execute('''INSERT INTO risks (user_id, domain, asset_name, threat, risk_level, analysis)
                    VALUES (?, ?, ?, ?, ?, ?)''',
                (session['user_id'], data.get('domain'), data.get('asset'),
                 data.get('threat'), 'HIGH', content))
    conn.commit()
    conn.close()
    
    return jsonify({'success': True, 'analysis': content})

@app.route('/api/generate-audit', methods=['POST'])
@login_required
def api_generate_audit():
    """Generate audit report."""
    lang = request.form.get('language', 'en')
    framework = request.form.get('framework', 'ISO 27001')
    audit_scope = request.form.get('audit_scope', 'full')
    domain = request.form.get('domain', 'Cyber Security')
    
    # Handle file upload
    evidence_files = request.files.getlist('evidence')
    evidence_info = []
    for f in evidence_files:
        if f and f.filename:
            evidence_info.append(f.filename)
    
    prompt = f"""Generate a comprehensive audit report for:
Framework: {framework}
Scope: {audit_scope}
Domain: {domain}
Evidence Documents: {', '.join(evidence_info) if evidence_info else 'No evidence provided'}

Include:
1. Executive Summary
2. Audit Scope & Methodology
3. Findings & Observations
4. Compliance Assessment
5. Risk Ratings
6. Recommendations
7. Action Plan"""

    if lang == 'ar':
        prompt = f"""أنشئ تقرير تدقيق شامل لـ:
الإطار: {framework}
النطاق: {audit_scope}
المجال: {domain}
وثائق الإثبات: {', '.join(evidence_info) if evidence_info else 'لم يتم تقديم أدلة'}

يتضمن:
1. الملخص التنفيذي
2. نطاق ومنهجية التدقيق
3. النتائج والملاحظات
4. تقييم الامتثال
5. تصنيف المخاطر
6. التوصيات
7. خطة العمل"""

    content = generate_ai_content(prompt, lang)
    return jsonify({'success': True, 'content': content})

@app.route('/api/generate-docx', methods=['POST'])
@login_required
def api_generate_docx():
    """Generate Word document from content."""
    from io import BytesIO
    
    data = request.json
    content = data.get('content', '')
    filename = data.get('filename', 'document')
    lang = data.get('language', 'en')
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Set RTL for Arabic
        if lang == 'ar':
            for section in doc.sections:
                section.page_width = Inches(8.5)
                section.page_height = Inches(11)
        
        # Add title
        title = doc.add_heading(filename.replace('_', ' ').title(), 0)
        if lang == 'ar':
            title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add content
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
            
            if line.startswith('## '):
                h = doc.add_heading(line[3:], level=1)
                if lang == 'ar':
                    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif line.startswith('### '):
                h = doc.add_heading(line[4:], level=2)
                if lang == 'ar':
                    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
                if lang == 'ar':
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p = doc.add_paragraph(line)
                if lang == 'ar':
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        from flask import send_file
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{filename}.docx'
        )
    except ImportError:
        # python-docx not installed, return error
        return jsonify({'error': 'Word generation not available'}), 500

@app.route('/api/set-language/<lang>')
def set_language(lang):
    """Set language preference."""
    session['lang'] = lang if lang in ['en', 'ar'] else 'en'
    return jsonify({'success': True, 'lang': session['lang']})

# ============================================================================
# MAIN
# ============================================================================

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
